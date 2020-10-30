
# import docx NOT python-docx
import docx

# create an instance of a word document
from docx import document

doc = docx.Document()

# add a heading of level 0 (largest heading)
doc.add_heading('Creating a document with python.', 0)

# add a paragraph and store
# the object in a variable
doc_para = doc.add_paragraph('Learning python the easy way, ')

# add a run i.e, style like
# bold, italic, underline, etc.
doc_para.add_run('you must practice regularly').bold = True
doc_para.add_run(', and ')
doc_para.add_run('study enough python library').italic = True

# add another paragraph and creating a list of topics in python
doc.add_heading('Key Topics You Must Know As An Intermediate Python Programmer')
doc_para = doc.add_paragraph('')
doc_para.add_run('1. Data Structure').italic = True
doc_para = doc.add_paragraph('')
doc_para.add_run('2. Function').italic = True
doc_para = doc.add_paragraph('')
doc_para.add_run('3. File').italic = True
doc_para = doc.add_paragraph('')
doc_para.add_run('4. Modules').italic = True
doc_para = doc.add_paragraph('')
doc_para.add_run('5. Object Oriented Programming').italic = True


# Get table data
table_items =(
    (1, 'Django', 'Python'),
    (2, '.NET', 'C#'),
)

# add a table
table = doc.add_table(1, 3)

# populate header row
header_cell = table.rows[0].cells
header_cell[0].text = 'SN'
header_cell[1].text = 'FRAMEWORK'
header_cell[2].text = 'LANGUAGE'

# add a data row for each item
for item in table_items:
    cells = table.add_row().cells
    cells[0].text = str(item.no)
    cells[1].text = item
    cells[2].text = item



# add a page break to start a new page
doc.add_page_break()

# add a heading of level 2
doc.add_heading('This python banner looks cool', 2)

# pictures can also be added to our word document
# width is optional
doc.add_picture('C:/Users/5th gen/Pictures/Saved Pictures/python images (1).jpeg')

# now save the document to a location
doc.save('C:/Users/5th gen/PycharmProjects/PythonTutorial/document/_output')